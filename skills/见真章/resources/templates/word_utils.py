"""
Potato skill — Word 文档生成工具。

用于将 potato 分析结果输出为格式正确的 .docx 文件。
解决 python-docx 对东亚字体支持不完整的问题：
  - 通过 XML 操作同步设置 w:ascii / w:hAnsi / w:eastAsia 三种字体
  - 提供标题、正文、加粗段落等预设样式
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from lxml import etree
import os


def _set_run_font(run, font_name, size=None, bold=None, color=None):
    """安全设置 run 的字体（同时覆盖 ascii / hAnsi / eastAsia）。

    python-docx 的 run.font.name 只设置 w:ascii，东亚字符会回退到
    默认字体导致显示混乱。此函数直接操作底层 XML 的 rFonts 元素。
    """
    rPr = run._r.get_or_add_rPr()

    # --- 字体名：ascii, hAnsi, eastAsia 三者同步 ---
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = etree.SubElement(rPr, qn('w:rFonts'))
    rFonts.set(qn('w:ascii'), font_name)
    rFonts.set(qn('w:hAnsi'), font_name)
    rFonts.set(qn('w:eastAsia'), font_name)
    # 也设置 cs (复杂脚本)
    rFonts.set(qn('w:cs'), font_name)

    if size is not None:
        run.font.size = size
    if bold is not None:
        run.bold = bold
    if color is not None:
        run.font.color.rgb = color


class PotatoDoc:
    """Potato 分析文档构建器。

    用法:
        doc = PotatoDoc('output.docx', body_font='宋体', title_font='微软雅黑')
        doc.add_title('当品牌想当你妈的第二个老公')
        doc.add_subtitle('OPPO 母亲节文案翻车背后的结构性困局')
        doc.add_body('这是正文段落……')
        doc.add_section('一、翻车绕开了产品，次次打在表达上')
        doc.save()
    """

    def __init__(self, output_path, body_font='宋体', title_font='微软雅黑',
                 body_size=11, line_spacing=1.6):
        self.doc = Document()
        self.output_path = output_path
        self.body_font = body_font
        self.title_font = title_font
        self.body_size = body_size
        self.line_spacing = line_spacing

        # 页面边距
        for section in self.doc.sections:
            section.top_margin = Cm(2.5)
            section.bottom_margin = Cm(2.5)
            section.left_margin = Cm(2.8)
            section.right_margin = Cm(2.8)

        # 全局默认段落样式：行距 + 首行无特殊格式
        style = self.doc.styles['Normal']
        style.paragraph_format.line_spacing = line_spacing
        style.paragraph_format.space_after = Pt(6)
        style.paragraph_format.first_line_indent = Cm(0)

    # ── 段落级别方法 ──────────────────────────────────

    def _make_run(self, paragraph, text, font_name, size, bold=False, color=None):
        """向 paragraph 追加一个 run 并设置好东亚字体。"""
        run = paragraph.add_run(text)
        _set_run_font(run, font_name, size=size, bold=bold, color=color)
        return run

    def add_title(self, text):
        """主标题（居中、大号、加粗）。"""
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(4)
        self._make_run(p, text, self.title_font, Pt(18), bold=True)

    def add_subtitle(self, text):
        """副标题（居中、小号、灰色）。"""
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(16)
        self._make_run(p, text, self.title_font, Pt(10),
                       color=RGBColor(0x66, 0x66, 0x66))

    def add_section(self, text):
        """节标题（左对齐、加粗、比正文字号大 2pt）。"""
        p = self.doc.add_paragraph()
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after = Pt(6)
        self._make_run(p, text, self.title_font, Pt(self.body_size + 2), bold=True)

    def add_body(self, text):
        """正文段落（首行缩进 0.7cm）。"""
        p = self.doc.add_paragraph()
        p.paragraph_format.first_line_indent = Cm(0.7)
        self._make_run(p, text, self.body_font, Pt(self.body_size))

    def add_blank(self, points=6):
        """空行（可指定间距）。"""
        p = self.doc.add_paragraph()
        p.paragraph_format.space_after = Pt(points)
        # 空 run 占位
        self._make_run(p, '', self.body_font, Pt(self.body_size))

    def add_separator(self, text='* * *'):
        """居中分隔符（三段式 * * *）。"""
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(12)
        self._make_run(p, text, self.body_font, Pt(10),
                       color=RGBColor(0x99, 0x99, 0x99))

    # ── 输出 ──────────────────────────────────────────

    def save(self):
        self.doc.save(self.output_path)
        print(f'[potato] Document saved → {self.output_path}')
