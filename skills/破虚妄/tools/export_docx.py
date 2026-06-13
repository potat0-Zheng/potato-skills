"""
事实核查报告 Word 导出工具
============================
将结构化报告导出为格式化的 .docx 文件。

支持两种调用方式：

1. 传统模式（固定9章节事实核查报告）：
   export_report(report_dict, output_dir=".")

2. 灵活模式（自定义章节列表，支持多种章节类型）：
   export_report_v2(report_dict, output_dir=".")

灵活模式支持的章节类型：
  - "heading": 独立标题
  - "text":   正文段落（支持多段，用 \\n 分隔）
  - "quote":  引用块（左缩进 + 灰色边框风格）
  - "table":  表格（需 headers + rows）
  - "timeline": 时间线（需 events 列表，每项含 date + text）
  - "keyvalue": 键值对列表

使用示例见模块末尾。
"""

from __future__ import annotations

import os
from datetime import datetime
from typing import Any, Optional, Union

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


# ---------------------------------------------------------------------------
# 公开函数
# ---------------------------------------------------------------------------


def export_report(
    report: dict[str, Any],
    output_dir: str = ".",
    filename: Optional[str] = None,
) -> str:
    """将结构化事实核查报告导出为 Word 文档（传统固定章节模式）。

    Args:
        report: 结构化报告数据。sections 为 dict（传统9章节）时走固定布局；
                 sections 为 list 时自动路由到 export_report_v2 灵活布局。
        output_dir: 输出目录。
        filename: 文件名（不含扩展名），默认从标题生成。

    Returns:
        生成的 .docx 文件的绝对路径。
    """
    sections = report.get("sections", {})
    if isinstance(sections, list):
        return export_report_v2(report, output_dir=output_dir, filename=filename)
    return _build_legacy(report, output_dir, filename)


def export_report_v2(
    report: dict[str, Any],
    output_dir: str = ".",
    filename: Optional[str] = None,
) -> str:
    """将结构化报告导出为 Word 文档（灵活自定义章节模式）。

    report["sections"] 为一个 list，每项为 dict，格式如下：

        {"type": "heading", "level": 1, "text": "章节标题"}
        {"type": "text", "text": "段落文本（支持 \\n 分段）"}
        {"type": "quote", "text": "引用内容", "attribution": "来源"}
        {"type": "table", "headers": ["列1", "列2"], "rows": [["a","b"]]}
        {"type": "timeline", "events": [{"date": "...", "text": "..."}]}
        {"type": "keyvalue", "pairs": [{"key": "...", "value": "..."}]}

    Returns:
        生成的 .docx 文件的绝对路径。
    """
    doc = _setup_document(report)
    custom_sections = report.get("sections", [])

    for sec in custom_sections:
        stype = sec.get("type", "text")
        if stype == "heading":
            _render_heading(doc, sec)
        elif stype == "text":
            _render_text(doc, sec)
        elif stype == "quote":
            _render_quote(doc, sec)
        elif stype == "table":
            _render_table(doc, sec)
        elif stype == "timeline":
            _render_timeline(doc, sec)
        elif stype == "keyvalue":
            _render_keyvalue(doc, sec)
        else:
            _render_text(doc, {"text": f"[未知章节类型: {stype}]"})

    _add_footer(doc)
    return _save_doc(doc, report, output_dir, filename)


# ---------------------------------------------------------------------------
# 内部：文档初始化 & 保存
# ---------------------------------------------------------------------------


def _setup_document(report: dict[str, Any]) -> Document:
    doc = Document()
    _configure_styles(doc)

    # 封面
    title = report.get("title", "未命名报告")
    h = doc.add_heading(title, level=0)
    _set_run_font(h, "微软雅黑")

    if url := report.get("source_url", ""):
        p = doc.add_paragraph()
        _add_styled_run(p, f"原文链接：{url}", size=Pt(9))

    if author := report.get("source_author", ""):
        p = doc.add_paragraph()
        _add_styled_run(p, f"原文作者：{author}", size=Pt(9))

    if sdate := report.get("source_date", ""):
        p = doc.add_paragraph()
        _add_styled_run(p, f"发布时间：{sdate}", size=Pt(9))

    return doc


def _save_doc(
    doc: Document,
    report: dict[str, Any],
    output_dir: str,
    filename: Optional[str],
) -> str:
    os.makedirs(output_dir, exist_ok=True)
    if filename is None:
        title = report.get("title", "报告")
        safe = "".join(c for c in title if c not in r'\/:*?"<>|')
        safe = safe[:40] if len(safe) > 40 else safe
        filename = safe
    output_path = os.path.join(output_dir, f"{filename}.docx")
    doc.save(output_path)
    return os.path.abspath(output_path)


def _add_footer(doc: Document) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _add_styled_run(p, "— 本报告由事实核查技能(fact-check)生成 —",
                    size=Pt(9), color=RGBColor(128, 128, 128))


# ---------------------------------------------------------------------------
# 内部：灵活章节渲染
# ---------------------------------------------------------------------------


def _render_heading(doc: Document, sec: dict[str, Any]) -> None:
    level = sec.get("level", 1)
    h = doc.add_heading(sec.get("text", ""), level=level)
    _set_run_font(h, "微软雅黑")


def _render_text(doc: Document, sec: dict[str, Any]) -> None:
    raw = sec.get("text", "")
    bold = sec.get("bold", False)
    for paragraph in raw.split("\n"):
        paragraph = paragraph.strip()
        if not paragraph:
            continue
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.first_line_indent = Cm(0.74)
        _add_styled_run(p, paragraph, bold=bold)


def _render_quote(doc: Document, sec: dict[str, Any]) -> None:
    text = sec.get("text", "")
    attribution = sec.get("attribution", "")
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1.5)
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(8)
    _add_styled_run(p, text, size=Pt(10), italic=True, color=RGBColor(80, 80, 80))
    if attribution:
        attr_p = doc.add_paragraph()
        attr_p.paragraph_format.left_indent = Cm(1.5)
        attr_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        _add_styled_run(attr_p, f"— {attribution}", size=Pt(9),
                        color=RGBColor(128, 128, 128))


def _render_table(doc: Document, sec: dict[str, Any]) -> None:
    headers = sec.get("headers", [])
    rows = sec.get("rows", [])
    if not headers:
        return
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Light Grid Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, hd in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = str(hd)
        for pp in cell.paragraphs:
            for r in pp.runs:
                r.bold = True
                r.font.size = Pt(10)
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            if ci < len(headers):
                cell = table.rows[ri + 1].cells[ci]
                cell.text = str(val)
                for pp in cell.paragraphs:
                    for r in pp.runs:
                        r.font.size = Pt(10)
    doc.add_paragraph()


def _render_timeline(doc: Document, sec: dict[str, Any]) -> None:
    events = sec.get("events", [])
    if not events:
        return
    for ev in events:
        date_str = ev.get("date", "")
        text = ev.get("text", "")
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        _add_styled_run(p, f"{date_str}  ", bold=True, size=Pt(10))
        _add_styled_run(p, text, size=Pt(10))
    doc.add_paragraph()


def _render_keyvalue(doc: Document, sec: dict[str, Any]) -> None:
    pairs = sec.get("pairs", [])
    if not pairs:
        return
    for item in pairs:
        key = item.get("key", "")
        value = item.get("value", "")
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        _add_styled_run(p, f"{key}：", bold=True, size=Pt(10.5))
        _add_styled_run(p, value, size=Pt(10.5))


# ---------------------------------------------------------------------------
# 内部：传统模式（向后兼容）
# ---------------------------------------------------------------------------


def _build_legacy(
    report: dict[str, Any],
    output_dir: str,
    filename: Optional[str],
) -> str:
    doc = _setup_document(report)
    sections = report.get("sections", {})

    _render_text(doc, {"text": sections.get("summary", ""), "bold": False})
    _legacy_claims(doc, sections.get("claims", []))
    _legacy_verdicts(doc, sections.get("verdicts", []))
    _legacy_evidence(doc, sections.get("evidence", []))
    _render_text(doc, {"text": sections.get("contradictions", ""), "bold": False})
    _legacy_opinions(doc, sections.get("opinion_analysis", []))
    _legacy_credibility(doc, sections.get("credibility", {}))
    _render_text(doc, {"text": sections.get("methodology", ""), "bold": False})
    _legacy_timestamp(doc, sections.get("timestamp", {}))

    _add_footer(doc)
    return _save_doc(doc, report, output_dir, filename)


def _legacy_claims(doc: Document, claims: list[dict[str, str]]) -> None:
    if not claims:
        return
    _render_heading(doc, {"level": 1, "text": "2. 事实断言清单"})
    _render_text(doc, {"text": "以下为从报道中拆解出的可核查事实断言与观点/定性的陈述："})
    _render_table(doc, {
        "headers": ["编号", "事实断言", "类型"],
        "rows": [[c.get("id", ""), c.get("text", ""), c.get("type", "")] for c in claims],
    })


def _legacy_verdicts(doc: Document, verdicts: list[dict[str, str]]) -> None:
    if not verdicts:
        return
    _render_heading(doc, {"level": 1, "text": "3. 核查结论"})
    _render_text(doc, {"text": "逐条标注：真实 / 部分真实 / 失实 / 证据不足，并简述理由。"})
    _render_table(doc, {
        "headers": ["编号", "结论", "理由"],
        "rows": [[v.get("id", ""), v.get("verdict", ""), v.get("reason", "")] for v in verdicts],
    })


def _legacy_evidence(doc: Document, evidence: list[dict[str, str]]) -> None:
    if not evidence:
        return
    _render_heading(doc, {"level": 1, "text": "4. 关键证据与来源"})
    _render_table(doc, {
        "headers": ["证据项", "来源", "链接/备注"],
        "rows": [[e.get("item", ""), e.get("source", ""), e.get("link", "")] for e in evidence],
    })


def _legacy_opinions(doc: Document, opinions: list[dict[str, str]]) -> None:
    if not opinions:
        return
    _render_heading(doc, {"level": 1, "text": "6. 观点辨析与推理说明"})
    for op in opinions:
        label = op.get("label", "")
        text = op.get("text", "")
        _render_text(doc, {"text": f'{label}："{text}"', "bold": True})
        if reasoning := op.get("reasoning", ""):
            _render_text(doc, {"text": reasoning})
        if conclusion := op.get("conclusion", ""):
            _render_text(doc, {"text": f"结论：{conclusion}"})


def _legacy_credibility(doc: Document, cred: dict[str, Any]) -> None:
    if not cred:
        return
    _render_heading(doc, {"level": 1, "text": "7. 综合可信度评估"})
    if truth := cred.get("truth_judgment", ""):
        _render_text(doc, {"text": truth})
    if sufficiency := cred.get("evidence_sufficiency", ""):
        _render_text(doc, {"text": sufficiency})
    if coverage := cred.get("coverage_completeness", ""):
        _render_text(doc, {"text": f"信息覆盖完整性：{coverage}"})
    if limitations := cred.get("limitations", []):
        _render_text(doc, {"text": "局限性说明：", "bold": True})
        for i, lim in enumerate(limitations, 1):
            _render_text(doc, {"text": f"{i}. {lim}"})
    if overall := cred.get("overall", ""):
        _render_text(doc, {"text": f"综合可信度：{overall}。"})


def _legacy_timestamp(doc: Document, ts: dict[str, str]) -> None:
    if not ts:
        return
    _render_heading(doc, {"level": 1, "text": "9. 核查时间戳"})
    _render_table(doc, {
        "headers": ["项目", "时间"],
        "rows": [
            ["原始报道发布时间", ts.get("source_published", "")],
            ["原始事件发生时间", ts.get("event_time", "")],
            ["本次核查时间", ts.get("report_time", "")],
        ],
    })


# ---------------------------------------------------------------------------
# 样式工具
# ---------------------------------------------------------------------------


def _configure_styles(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)
    style = doc.styles["Normal"]
    style.font.name = "宋体"
    style.font.size = Pt(10.5)
    style.element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")


def _set_run_font(paragraph_obj, font_name: str) -> None:
    for run in paragraph_obj.runs:
        run.font.name = font_name
        run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)


def _add_styled_run(
    para,
    text: str,
    bold: bool = False,
    italic: bool = False,
    size=None,
    color=None,
) -> None:
    run = para.add_run(text)
    run.font.name = "宋体"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    if size is not None:
        run.font.size = size
    else:
        run.font.size = Pt(10.5)
    run.bold = bold
    run.italic = italic
    if color is not None:
        run.font.color.rgb = color


# ---------------------------------------------------------------------------
# 使用示例
# ---------------------------------------------------------------------------

if __name__ == "__main__":
    # ---- 示例1：灵活模式（编译型报告） ----
    report_v2 = {
        "title": "OPPO母亲节文案事件——武汉大学声明与媒体反应汇编",
        "source_url": "",
        "source_date": "2026-05-08",
        "sections": [
            {"type": "heading", "level": 1, "text": "一、事件概述"},
            {"type": "text", "text": "2026年5月8日，OPPO发布母亲节活动文案引发巨大争议……"},
            {"type": "heading", "level": 1, "text": "二、武汉大学声明"},
            {"type": "quote", "text": "极感诧异和震惊……极不认同此文案之内容……",
             "attribution": "武汉大学文学院（2026-05-10）"},
            {"type": "heading", "level": 1, "text": "三、媒体报道一览"},
            {"type": "table",
             "headers": ["媒体名称", "类型", "核心态度"],
             "rows": [
                 ["澎湃新闻", "官方媒体", "批评武大迎合\"舆论连坐\""],
                 ["新京报", "官方媒体", "武大没必要\"自我连坐\""],
             ]},
            {"type": "heading", "level": 1, "text": "四、事件时间线"},
            {"type": "timeline",
             "events": [
                 {"date": "5月8日", "text": "OPPO发布争议文案并致歉"},
                 {"date": "5月10日", "text": "武汉大学发布声明"},
                 {"date": "5月11日", "text": "OPPO再次致歉，高管被降级"},
             ]},
            {"type": "heading", "level": 1, "text": "五、关键信息"},
            {"type": "keyvalue",
             "pairs": [
                 {"key": "核查时间", "value": "2026-05-13"},
                 {"key": "信息来源数", "value": "14个独立来源"},
                 {"key": "综合可信度", "value": "高"},
             ]},
        ],
    }
    path = export_report_v2(report_v2, output_dir=".")
    print(f"灵活模式示例: {path}")

    # ---- 示例2：传统模式（向后兼容） ----
    report_v1 = {
        "title": "传统事实核查报告示例",
        "source_url": "https://example.com/article",
        "sections": {
            "summary": "原始内容精要……",
            "claims": [{"id": "F1", "text": "事实断言", "type": "可核查事实"}],
            "verdicts": [{"id": "F1", "verdict": "真实", "reason": "多源一致"}],
            "evidence": [{"item": "证据", "source": "来源", "link": "https://..."}],
            "contradictions": "无明显矛盾",
            "opinion_analysis": [{"label": "V1", "text": "观点", "reasoning": "分析", "conclusion": "结论"}],
            "credibility": {"truth_judgment": "高度吻合", "evidence_sufficiency": "充足", "overall": "高"},
            "methodology": "基于多源检索……",
            "timestamp": {"report_time": "2026-05-13", "event_time": "2026-05-08", "source_published": "2026-05-08"},
        },
    }
    path2 = export_report(report_v1, output_dir=".")
    print(f"传统模式示例: {path2}")
